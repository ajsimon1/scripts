"""
Dr lin gave dx codes in a word dox, using python docx to pull
them templates out and into an excel spreadsheet
"""

import docx
import openpyxl as xl

# set documents as Document object
# doc = docx.Document('GIMACROSPrasadedited7.as.docx')
# new file name for spreadsheet
dest_file = 'AGA_linDxTemps.xlsx'
# load workbook && active sheet
wb = xl.load_workbook(dest_file)
ws = wb.active
# load range in var
a_range = 'A2:B306'
b_range = 'D7:E306'
count = 1

"""
# pulled out the rows but couldnt get it formatted into xl right
# just found a way to copy and paste
tbls = doc.tables
for tble in tbls:
    for row_doc in tble.rows[0:2]:
        print('step 1')
        for cell in row_doc.cells:
            print('step 2')
            for paragraph in cell.paragraphs[0:2]:
                print('step 3')
                for row_xl in ws.iter_rows(a_range):
                    print('step 4')
                    for cell in row_xl:
                        print('step 5')
                        cell.value = paragraph.text
                        print('cell value is {cellvalue}, paragrapg text is {para}'.format(cellvalue=cell.value, para=paragraph.text))
                        continue
    # wb.save(dest_file)
"""

# add body site based on first character of column B
def add_site(a_range):
    for row in ws.iter_rows(a_range):
        if row[0].value:
            continue
        else:
            cell_text = row[1].value
            if cell_text[0:1] == 'E':
                row[0].value = 'Esophagus'
            elif cell_text[0:1] == 'S':
                row[0].value = 'Stomach'
            elif cell_text[0:1] == 'C':
                row[0].value = 'Colon'
            elif cell_text[0:1] == 'D':
                row[0].value = 'Duodenum'
            else:
                row[0].value = 'SITE NEEDED'
    wb.save(dest_file)


def split_micros(a_range):
    for row in ws.iter_rows(a_range):
        try:
            pattern = 'COMMENT:'
            s = row[0].value
            comment = s.split(pattern, 1)[1]
            row[1].value = 'COMMENT: ' + comment
            row[0].value = s.split(pattern, 1)[0]
        except IndexError:
            pass
    wb.save(dest_file)

# add_site(a_range)
split_micros(b_range)
